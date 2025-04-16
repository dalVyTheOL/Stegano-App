import sys
from PyQt5 import QtWidgets, uic
from PyQt5.QtWidgets import QFileDialog, QVBoxLayout
from docx import Document
from docx.shared import RGBColor
from PyQt5.QtCore import Qt

def extract_secret_message(input_file):
    """
    Функция извлекает секретное сообщение из текста с символами ударения.
    """
    doc = Document(input_file)
    secret_message = ""
    accent_char = "́"  # Unicode: U+0301 - символ ударения
    for paragraph in doc.paragraphs:
        for char_index in range(len(paragraph.text)):
            char = paragraph.text[char_index]
            if char == accent_char and char_index > 0:
                # Проверяем, что предыдущий символ - буква
                if paragraph.text[char_index - 1].isalpha():
                    secret_message += paragraph.text[char_index - 1]
    return secret_message

def extract_hidden_message(input_file):
    """
    Функция извлекает скрытое сообщение из текста с символами почти чёрного цвета.
    """
    doc = Document(input_file)
    hidden_message = ""
    hidden_message_color = (1, 0, 0)  # RGB-код почти чёрного цвета
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Проверяем цвет текущего символа
            if run.font.color.rgb == hidden_message_color:
                # Добавляем символ к найденному сообщению
                hidden_message += run.text
    return hidden_message

def hide_message_in_word(input_file, output_file, secret_message):
    """
    Функция скрывает сообщение в документе Word, окрашивая его символы в почти чёрный цвет.
    """
    doc = Document(input_file)
    secret_chars = list(secret_message)
    hidden_message_color = RGBColor(1, 0, 0)
    index = 0
    for paragraph in doc.paragraphs:
        runs = paragraph.runs
        paragraph.clear()  # Очищаем параграф от всех текстовых данных
        for run in runs:
            for char in run.text:
                if index < len(secret_chars) and char == secret_chars[index]:
                    new_run = paragraph.add_run(char)
                    new_run.font.color.rgb = hidden_message_color
                    index += 1
                else:
                    paragraph.add_run(char)
        if index == len(secret_chars):
            break
    doc.save(output_file)
    print("Секретное сообщение успешно скрыто в файле", output_file)


def add_accent_marks(text, secret_message):
    """
    Функция добавляет символы ударения над буквами в текст на основе заданного секретного сообщения.
    """
    modified_text = ""
    secret_index = 0
    for i, char in enumerate(text):
        if secret_index < len(secret_message) and char.isalpha():
            if secret_message[secret_index].lower() == char.lower():
                char += "́"  # Добавляем символ ударения
                secret_index += 1
        modified_text += char

        # Проверяем, достигли ли мы конца секретного сообщения
        if secret_index == len(secret_message):
            modified_text += text[i+1:]  # Добавляем оставшийся текст
            break
    secret_message = secret_message[secret_index:]

    return modified_text, secret_message

def write_text_with_accent_marks(input_file, output_file, secret_message):
    """
    Функция читает текст из файла Word, добавляет символы ударения в соответствии с секретным сообщением
    и записывает измененный текст в новый файл Word.
    """
    doc = Document(input_file)
    modified_doc = Document()
    for paragraph in doc.paragraphs:
        modified_text, secret_message = add_accent_marks(paragraph.text, secret_message)
        modified_doc.add_paragraph(modified_text)
    modified_doc.save(output_file)
    print("Символы ударения успешно добавлены в файл", output_file)


class InfoAuthor(QtWidgets.QDialog):
    def __init__(self):
        super(InfoAuthor, self).__init__()
        self.setWindowTitle("Об авторе")
        self.setFixedSize(370, 200)  # Установка фиксированного размера
        layout = QVBoxLayout()
        text_edit = QtWidgets.QTextEdit()
        text_edit.setReadOnly(True)  # Запретить редактирование текста
        text_edit.setText("Данное приложение спроектировал Левшенко Владислав Александрович, студент группы ИТ-1035222, "
                          "обучающийся на втором курсе специальности 'Информационная Безопасность Автоматизированных Систем'."
                          "\n"
                          "Приложение использовалось для защиты курсовой работы по дисциплине 'Методы и Средства Криптографической Защиты "
                          "Информации'.")
        layout.addWidget(text_edit)
        self.setLayout(layout)
class InfoDialog(QtWidgets.QDialog):
    def __init__(self):
        super(InfoDialog, self).__init__()
        self.setWindowTitle("Описание методов стеганографии")
        self.setFixedSize(550, 500)  # Установка фиксированного размера
        layout = QVBoxLayout()
        text_edit = QtWidgets.QTextEdit()
        text_edit.setReadOnly(True)  # Запретить редактирование текста
        text_edit.setText("СПРАВКА ОБ ИСПОЛЬЗУЕМЫХ МЕТОДАХ СТЕГАНОГРАФИИ"
                          "\n"
                          "\nМетод Микроточек:"
                          "\nДанный метод использовался германскими шпионами во времена Первой мировой войны. "
                          "\nЕго суть заключается в добавлении в текст маленьких точек, расположенных непосредственно над символами секретного сообщения."
                          "\nВ данной программе используются знаки ударения, которые находятся прямо над буквами в Word-файле."
                          "\nПрограмма получает на вход файл Word(.docx) и использует текст, который введёт пользователь для последующего скрытия в новом файле."
                          "\n"
                          "\nМетод изменения значения RGB:"
                          "\nДанный метод заключается в изменении цвета символов скрываемого сообщения."
                          "\nСтандартный цвет текста, вводимого в Word файле равен значениям RGB(0,0,0). Соответственно, чтобы скрыть сообщение "
                          "нужно слегка поменять значения RGB у символов секретного сообщения."
                          "\nВ данной программе изменение происходит с разницей всего в 1 значение красного(R). На выходе программа придаёт "
                          "символам, входящим в скрываемое сообщение значения RGB(1,0,0) или 'почти чёрный'. "
                          "Благодаря этому методу, пользователь при чтении текста не заметит разницу между обычным символом текста и символом "
                          "секретного сообщения.")
        layout.addWidget(text_edit)
        self.setLayout(layout)

class InfoUser(QtWidgets.QDialog):
    def __init__(self):
        super(InfoUser, self).__init__()
        self.setWindowTitle("Руководство для пользователя")
        self.setFixedSize(500, 500)  # Установка фиксированного размера
        layout = QVBoxLayout()
        text_edit = QtWidgets.QTextEdit()
        text_edit.setReadOnly(True)  # Запретить редактирование текста
        text_edit.setText("ПОШАГОВАЯ ИНСТРУКЦИЯ:"
                          "\n"
                          "\n"
                          "Первым делом выберите метод скрытия/раскрытия информации (Микроточки/RGB)"
                          "\n"
                          "\nДля скрытия:"
                          "\nПерейдите в область 'Скрытие информации' и выберите файл формата (.docx) с помощью кнопки или меню сверху. "
                          "\nЗатем введите в соответствующее поле информацию, которую хотите скрыть."
                          "\nПосле этого введите путь и название для нового Word-файла(.docx), в который будет записан текст со скрытой информацией."
                          "\nВ конце нажмите кнопку 'Скрыть', чтобы получить файл по указанному вами пути."
                          "\n"
                          "\nДля раскрытия:"
                          "\nПерейдите в область 'Раскрытие информации' и выберите файл формата (.docx) с помощью кнопки или меню сверху. "
                          "\nДалее нажмите на кнопку 'Раскрыть', после чего в соответствующем поле появиться скрытая информация."
                          "\n"
                          "\nПРИМЕЧАНИЯ:"
                          "\nДля более безопасного и незаметного скрытия информации рекомендуется использовать исходный файл "
                          "который содержит в себе как МИНИМУМ 1000 символов."
                          "\nЧтобы скрыть сообщение методом 'Микроточек' необходимо ввести секретное сообщение БЕЗ ПРОБЕЛОВ, "
                          "чтобы сообщение содержало только буквы.")
        layout.addWidget(text_edit)
        self.setLayout(layout)

class MyMainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super(MyMainWindow, self).__init__()
        # Загрузка формы
        uic.loadUi('kursovaya.ui', self)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowMaximizeButtonHint)
        self.butChoiceFind.clicked.connect(self.choose_file_find)
        self.butFind.clicked.connect(self.find_message)
        self.butChoiceHide.clicked.connect(self.choose_file_hide)
        self.butHide.clicked.connect(self.hide_message)
        self.for_in.triggered.connect(self.choose_file_hide)
        self.for_out.triggered.connect(self.choose_file_find)
        self.for_exit.triggered.connect(self.exit_program)
        self.buttonInfo.clicked.connect(self.show_info_dialog)
        self.help_list.triggered.connect(self.show_info_user)
        self.about_author.triggered.connect(self.show_info_author)

    def exit_program(self):
        QtWidgets.QApplication.quit()

    def choose_file_find(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        file_filter = "Документы Word (*.docx)"
        filename, _ = QFileDialog.getOpenFileName(self, "Выберите файл", "", file_filter, options=options)
        if filename:
            self.textSucc3.setText("Успешно!")  # Отображаем "Успешно!"
            self.input_file_find = filename  # Сохраняем выбранный файл для последующего использования

    def choose_file_hide(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        file_filter = "Документы Word (*.docx)"
        filename, _ = QFileDialog.getOpenFileName(self, "Выберите файл", "", file_filter, options=options)
        if filename:
            self.textSucc1.setText("Успешно!")  # Отображаем "Успешно!"
            self.input_file_hide = filename  # Сохраняем выбранный файл для последующего использования

    def find_message(self):
        if self.radButRGB.isChecked() and hasattr(self,
                                                  'input_file_find'):  # Проверяем, выбран ли radButRGB и выбран ли файл
            message = extract_hidden_message(self.input_file_find)
            if message:
                self.textOutFind.setText("{}".format(message))
            else:
                self.textOutFind.setText("Секретное сообщение не обнаружено")
        elif self.radButMicro.isChecked() and hasattr(self,
                                                      'input_file_find'):  # Проверяем, выбран ли radButMicro и выбран ли файл
            secret_message = extract_secret_message(self.input_file_find)
            if secret_message:
                self.textOutFind.setText("{}".format(secret_message))
            else:
                self.textOutFind.setText("Секретное сообщение не обнаружено")

    def hide_message(self):
        if self.radButRGB.isChecked() and hasattr(self,
                                                  'input_file_hide') and self.textInHide.toPlainText() and self.textPathNew.text():
            input_file = self.input_file_hide
            output_file = self.textPathNew.text()
            if not output_file.endswith('.docx'):
                self.textSucc2.setText("Неверный формат!")
                return
            secret_message = self.textInHide.toPlainText()
            hide_message_in_word(input_file, output_file, secret_message)
            self.textSucc2.setText("Успешно!")
        elif self.radButMicro.isChecked() and hasattr(self,
                                                      'input_file_hide') and self.textInHide.toPlainText() and self.textPathNew.text():
            input_file = self.input_file_hide
            output_file = self.textPathNew.text()
            if not output_file.endswith('.docx'):
                self.textSucc2.setText("Неверный формат!")
                return
            secret_message = self.textInHide.toPlainText()
            write_text_with_accent_marks(input_file, output_file, secret_message)
            self.textSucc2.setText("Успешно!")

    def show_info_dialog(self):  # Метод для отображения диалогового окна справки
        dialog = InfoDialog()
        dialog.exec_()

    def show_info_user(self):  # Метод для отображения диалогового окна справки
        dialog = InfoUser()
        dialog.exec_()

    def show_info_author(self):  # Метод для отображения диалогового окна справки
        dialog = InfoAuthor()
        dialog.exec_()

if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    window = MyMainWindow()
    window.show()
    sys.exit(app.exec_())
